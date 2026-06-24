"""
把 演讲稿.md 转成 演讲稿.docx
- 中文字体 Microsoft YaHei，西文 Arial
- 一级标题（# 答辩演讲稿）作为文档大标题
- 二级标题（## 第 N 页｜xxx）作为各页大节
- 引用块（>）灰色斜体
- 表格转为真表格
- ** 加粗 ** 渲染为粗体
- 段落正文 14pt（接近 PPT 演讲稿易读字号）
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path("/Users/rexanyang/Desktop/bishe/graduatejob-main/docs/ppt_v2/演讲稿.md")
DST = Path("/Users/rexanyang/Desktop/bishe/graduatejob-main/docs/ppt_v2/演讲稿.docx")


def set_run_chinese_font(run, font_name="Microsoft YaHei"):
    """让 run 同时使用中英文字体"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def add_runs_with_inline(paragraph, text, base_size=14, bold=False, italic=False, color=None):
    """处理 **粗体**、`代码` 这类内联，分段塞入 run"""
    # 同时拆 ** ** 和 ` `
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            set_run_chinese_font(run, "Consolas")
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            continue  # 已经设置过字体了，跳过下面通用设置
        else:
            run.text = part
            run.bold = bold
        run.italic = italic
        run.font.size = Pt(base_size)
        if color is not None:
            run.font.color.rgb = color
        set_run_chinese_font(run)


def add_paragraph_text(doc, text, base_size=14, bold=False, italic=False, color=None,
                       space_before=4, space_after=4, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if alignment is not None:
        p.alignment = alignment
    add_runs_with_inline(p, text, base_size=base_size, bold=bold, italic=italic, color=color)
    return p


def add_heading_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x82)
    set_run_chinese_font(run)


def add_heading_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    # 顶部加细边框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F4C82")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x82)
    set_run_chinese_font(run)


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    add_runs_with_inline(p, text, base_size=12, italic=True, color=RGBColor(0x64, 0x74, 0x8B))


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B5C7DD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table_from_rows(doc, rows):
    """rows: list[list[str]]，第 0 行视为表头"""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            # 清空默认段
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_with_inline(p, cell_text, base_size=11,
                                 bold=(r_idx == 0))
    doc.add_paragraph()  # 表格后留一段空


def parse_and_render(md_text: str) -> Document:
    doc = Document()

    # 全局默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:cs"), "Microsoft YaHei")

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # 分隔线
        if line.strip() == "---":
            add_hr(doc)
            i += 1
            continue

        # 一级标题
        m = re.match(r"^# (.+)$", line)
        if m:
            add_heading_h1(doc, m.group(1).strip())
            i += 1
            continue

        # 二级标题
        m = re.match(r"^## (.+)$", line)
        if m:
            add_heading_h2(doc, m.group(1).strip())
            i += 1
            continue

        # 三级标题
        m = re.match(r"^### (.+)$", line)
        if m:
            add_paragraph_text(doc, m.group(1).strip(),
                               base_size=15, bold=True,
                               color=RGBColor(0x1F, 0x49, 0x7D),
                               space_before=10, space_after=4)
            i += 1
            continue

        # 引用块（连续 > 行合成一段）
        if line.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip("> ").rstrip())
                i += 1
            add_quote(doc, "\n".join(buf))
            continue

        # 表格（| 开头，且下一行是分隔行）
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-:\s|]+\|$", lines[i + 1].strip()):
            rows = []
            # header
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header)
            i += 2  # 跳过表头和分隔行
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            add_table_from_rows(doc, rows)
            continue

        # 普通段落
        add_paragraph_text(doc, line.strip(), base_size=14)
        i += 1

    return doc


def main():
    md_text = SRC.read_text(encoding="utf-8")
    doc = parse_and_render(md_text)
    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"✅ 生成成功: {DST}  size={DST.stat().st_size} bytes")


if __name__ == "__main__":
    main()
