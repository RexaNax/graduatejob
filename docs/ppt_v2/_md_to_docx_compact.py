"""
把 演讲稿_精简版.md 转成紧凑两页 docx
- A4 + 窄边距(1.5cm)
- 正文 10pt / 1.15 行距 / 段后 2pt
- 标题与正文同一行（**Pn xxx｜时长** 直接当段首加粗），不另起标题
- 中文 Microsoft YaHei
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path("/Users/rexanyang/Desktop/bishe/graduatejob-main/docs/ppt_v2/演讲稿_精简版.md")
DST = Path("/Users/rexanyang/Desktop/bishe/graduatejob-main/docs/ppt_v2/演讲稿_精简版.docx")


def set_run_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def add_runs(paragraph, text, base_size=10, color=None):
    """处理 **bold** 和 `code`"""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
            set_run_font(run)
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            set_run_font(run, "Consolas")
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            continue
        else:
            run.text = part
            set_run_font(run)
        run.font.size = Pt(base_size)
        if color is not None:
            run.font.color.rgb = color


def add_para(doc, text, size=10, bold=False, color=None,
             before=0, after=2, line_spacing=1.15, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    add_runs(p, text, base_size=size, color=color)
    if bold:
        for r in p.runs:
            r.bold = True
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x82)
    set_run_font(run)


def add_hr_thin(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B5C7DD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    doc = Document()

    # 默认样式
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:cs"), "Microsoft YaHei")

    # A4 + 窄边距
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    md_text = SRC.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            add_hr_thin(doc)
            i += 1
            continue

        # 一级标题（# xxx）
        m = re.match(r"^# (.+)$", line)
        if m:
            add_title(doc, m.group(1).strip())
            i += 1
            continue

        # 引用块
        if line.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip("> ").rstrip())
                i += 1
            add_para(doc, " ".join(buf), size=9,
                     color=RGBColor(0x64, 0x74, 0x8B),
                     before=0, after=4, line_spacing=1.2)
            continue

        # 正文段
        add_para(doc, line.strip(), size=10, before=0, after=2, line_spacing=1.18)
        i += 1

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"✅ 生成成功: {DST}  size={DST.stat().st_size} bytes")


if __name__ == "__main__":
    main()
