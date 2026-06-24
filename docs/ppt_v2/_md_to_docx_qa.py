"""
把 老师可能问的问题.md 转成紧凑 docx
- A4 窄边距
- 正文 10pt / 1.18 行距
- 章节标题 (## 一、xxx) 蓝色粗体带下划线
- 题目 (### Qx. xxx) 深蓝粗体
- "3秒抓手/展开回答/追问兜底" 加底色突出
- 有序列表正确缩进
- 表格转真表格
- 中文 Microsoft YaHei
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path("/Users/rexanyang/Desktop/bishe/graduatejob-main/docs/ppt_v2/老师可能问的问题.md")
DST = Path("/Users/rexanyang/Desktop/bishe/graduatejob-main/docs/ppt_v2/老师可能问的问题.docx")


def set_run_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


# 三个关键标记词 → 高亮颜色（背景灰，字深蓝）
HIGHLIGHT_KEYWORDS = ["3 秒抓手", "展开回答", "追问兜底", "关键事实", "比喻"]


def add_runs(paragraph, text, base_size=10, color=None, bold_default=False):
    """处理 **bold**、`code`、以及关键词高亮"""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            run.font.size = Pt(base_size)
            # 是否是关键词，给个色
            if any(kw in inner for kw in HIGHLIGHT_KEYWORDS):
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x4A)  # 深红
            elif color is not None:
                run.font.color.rgb = color
            else:
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
            set_run_font(run)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas")
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = paragraph.add_run(part)
            run.bold = bold_default
            run.font.size = Pt(base_size)
            if color is not None:
                run.font.color.rgb = color
            set_run_font(run)


def add_para(doc, text, size=10, before=0, after=2, line_spacing=1.18,
             color=None, alignment=None, bold=False, indent_left_cm=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    if indent_left_cm is not None:
        pf.left_indent = Cm(indent_left_cm)
    if alignment is not None:
        p.alignment = alignment
    add_runs(p, text, base_size=size, color=color, bold_default=bold)
    return p


def add_doc_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x82)
    set_run_font(run)


def add_section_h2(doc, text):
    """## 一、xxx —— 章节大标题"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    # 上下边框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for tag in ("top", "bottom"):
        b = OxmlElement(f"w:{tag}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), "0F4C82")
        pBdr.append(b)
    pPr.append(pBdr)
    # 浅蓝底
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E6F0FA")
    pPr.append(shd)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x82)
    set_run_font(run)


def add_question_h3(doc, text):
    """### Qx. xxx —— 题目"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(7)
    pf.space_after = Pt(2)
    pf.keep_with_next = True  # 题目跟下一段一起，不被分页拆开
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    set_run_font(run)


def add_hr_thin(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B5C7DD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_quote(doc, text):
    p = add_para(doc, text, size=9.5,
                 before=2, after=4, line_spacing=1.25,
                 color=RGBColor(0x55, 0x66, 0x78),
                 indent_left_cm=0.4)
    # 左侧加竖条
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B5C7DD")
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def add_table_from_rows(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_runs(p, cell_text, base_size=9.5,
                     bold_default=(r_idx == 0))
            if r_idx == 0:
                for r in p.runs:
                    r.bold = True
    # 表格后留半行
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(2)


def main():
    doc = Document()

    # 默认样式
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:cs"), "Microsoft YaHei")

    # A4 窄边距
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

    md_text = SRC.read_text(encoding="utf-8")
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 跳空行
        if not line.strip():
            i += 1
            continue

        # ---
        if line.strip() == "---":
            add_hr_thin(doc)
            i += 1
            continue

        # 一级标题
        m = re.match(r"^# (.+)$", line)
        if m:
            add_doc_title(doc, m.group(1).strip())
            i += 1
            continue

        # 二级标题
        m = re.match(r"^## (.+)$", line)
        if m:
            add_section_h2(doc, m.group(1).strip())
            i += 1
            continue

        # 三级标题
        m = re.match(r"^### (.+)$", line)
        if m:
            add_question_h3(doc, m.group(1).strip())
            i += 1
            continue

        # 引用块
        if line.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip("> ").rstrip())
                i += 1
            add_quote(doc, " ".join(buf))
            continue

        # 表格
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-:\s|]+\|$", lines[i + 1].strip()):
            rows = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            add_table_from_rows(doc, rows)
            continue

        # 有序列表 1. 2. 3.
        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            text = f"{m.group(1)}. {m.group(2)}"
            add_para(doc, text, size=10, before=0, after=2,
                     line_spacing=1.18, indent_left_cm=0.6)
            i += 1
            continue

        # 无序列表 - 开头
        m = re.match(r"^-\s+(.+)$", line)
        if m:
            text = f"• {m.group(1)}"
            add_para(doc, text, size=10, before=0, after=2,
                     line_spacing=1.18, indent_left_cm=0.6)
            i += 1
            continue

        # 普通段落
        add_para(doc, line.strip(), size=10, before=0, after=2, line_spacing=1.18)
        i += 1

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"✅ 生成成功: {DST}")
    print(f"   大小: {DST.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
